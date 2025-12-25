import pandas as pd
from docxtpl import DocxTemplate
from docxcompose.composer import Composer
from docx import Document
import os

# --- Configurações Iniciais ---#
NOME_ARQUIVO_EXCEL = "benchmark.xlsx"
NOME_ABA_EXCEL = "Matriz_Aceitacao"
NOME_ARQUIVO_MODELO = "arquivo_final.docx"
NOME_ARQUIVO_SAIDA = "Relatorio_Final_Completo.docx"

print("=== INICIANDO AUTOMAÇÂO ===")

# 1. Carregar os dados do Excel
print(f"Lendo arquivo Excel: {NOME_ARQUIVO_EXCEL}...")
try:
    df = pd.read_excel(NOME_ARQUIVO_EXCEL, sheet_name=NOME_ABA_EXCEL)
    print(f"Sucesso! Encontrei {len(df)} empresas para processar.")
except Exception as e:
    print(f"ERRO FATAL ao ler Excel: {e}")
    exit()

# Função para formatar a moeda (R$ 1.000,00)
def formatar_moeda(valor):
    try:
        texto = f"{float(valor):,.f}"
        return "R$ " + texto.replace(",", "X").replace(".", ",").replace("X",".")
    except:
        return valor # Se não for número, devolve como está
    
# Lista para guardar os nomes dos arquivos temporários temporários de cada empresa
arquivos_temporarios = []


# 2. O Loop Principal (Para cada linha do Excel)
print("Começando a gerar os documentos individuais...")

for index, linha in df.iterrows():
    # a) Carrega o modelo original zerado
    doc = DocxTemplate(NOME_ARQUIVO_MODELO)
    
    # b) Pega os dados da linha atual e organiza para o Word
    # Atenção: Aqui ligamos o Excel (esquerda com as Tags do Word (direita)
    contexto = {
        'nome_empresa': linha['Nome da Empresa'],
        'atividade':    linha['Atividade da Empresa'],
        'funcionarios': linha['Funcionários'],
        'gasto_anual': formatar_moeda(linha['Gasto Anual']),    # Formatando a moeda
        'faturamento': formatar_moeda(linha['Faturamento Anual']),    # Formatando a moeda
    }

    # c) Substitui as tags pelos dados reais
    doc.render(contexto)
    
    # d) Salva um arquivo temporário só dessa empresa
    nome_temp = f"temp_empresa_{index}.docx"
    doc.save(nome_temp)
    arquivos_temporarios.append(nome_temp)
    
    print(f" -> Processado: {linha['Nome da Empresa']}")
    
    
# 3. A Fusão (Juntar arquivos individuais num unico arquivo)
print("Unindo todos os arquivos em um só...")

if arquivos_temporarios:
    # Pega o primeiro arquivo para ser a base (mestre)
    master_doc = Document(arquivos_temporarios[0])
    composer = Composer(master_doc)
    
    # Adiciona os outros arquivos um por um
    for arq_temp in arquivos_temporarios[1:]:
        doc_to_append = Document(arq_temp)
    
        # Adiciona uma quebra de linha/página antes de colar o próximo
        master_doc.add_page_break()
        # DICA: Se quiser apenas espaço de linha e não pular folha,
        # troque a linha acima por: master_doc.add_paragraph('\n')
        
        composer.append(doc_to_append)


# 4. Salvar o Arquivo Final
    composer.save(NOME_ARQUIVO_SAIDA)
    print(f"\n=== SUCESSO! Arquivo salvo como: {NOME_ARQUIVO_SAIDA} ===")


# 5. Limpeza (Apagar os temporários)
    print("Limpando arquivos temporários...")
    for f in arquivos_temporarios:
        try:
            os.remove(f)
        except:
            pass
    print("Limpeza concluída")
else:
    print("Nenhum dado foi processo")





