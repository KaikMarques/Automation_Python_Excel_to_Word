from pathlib import Path
from tkinter import Tk, Canvas, PhotoImage, filedialog, messagebox
import pandas as pd
from docxtpl import DocxTemplate
from docxcompose.composer import Composer
from docx import Document
import os

# --- CONFIGURAÇÕES INICIAIS ---
OUTPUT_PATH = Path(__file__).parent
ASSETS_PATH = OUTPUT_PATH / Path(r"C:\Users\kaikm\OneDrive\Documentos\Projetos\Automação\Automation_Python_Excel_to_Word\design\build\assets\frame0")

def relative_to_assets(path: str) -> Path:
    return ASSETS_PATH / Path(path)

# --- VARIÁVEIS GLOBAIS ---
caminho_excel = ""
caminho_modelo_word = ""
pasta_destino = ""

# Dicionário para guardar os IDs dos textos na tela
elementos_texto = {
    "excel": None,
    "word": None,
    "pasta": None
}

# --- FUNÇÕES DE AUXÍLIO ---

def encurtar_caminho(caminho):
    """Pega um caminho longo e retorna só o final"""
    if not caminho:
        return ""
    nome = os.path.basename(caminho)
    if len(nome) > 30:
        return "..." + nome[-27:]
    return nome

def atualizar_label(tipo, texto, cor):
    """Atualiza o texto desenhado no canvas"""
    # Se já existir um texto criado pelo Python, deleta ele antes de escrever o novo
    if elementos_texto[tipo]:
        canvas.delete(elementos_texto[tipo])
    
    # Se o texto for vazio (caso de limpar), para por aqui e não escreve nada
    if not texto:
        return

    # Define as posições Y baseadas nos botões
    posicoes_y = {
        "excel": 222,
        "word": 291,
        "pasta": 360
    }
    
    # Cria o novo texto (apenas quando selecionar arquivo)
    id_novo = canvas.create_text(
        430.0,              # X
        posicoes_y[tipo],   # Y
        text=texto,
        anchor="e",         # Alinhado à direita
        fill=cor,
        font=("Arial", 10, "bold")
    )
    elementos_texto[tipo] = id_novo

# --- FUNÇÕES DOS BOTÕES (LÓGICA) ---

def gerenciar_excel(evento_limpar=False):
    global caminho_excel
    
    if evento_limpar:
        caminho_excel = ""
        atualizar_label("excel", None, None) # Limpa o texto da tela
        return

    arquivo = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx")])
    if arquivo:
        caminho_excel = arquivo
        # Só escreve agora que selecionou
        atualizar_label("excel", encurtar_caminho(arquivo), "#4CAF50") 

def gerenciar_word(evento_limpar=False):
    global caminho_modelo_word
    
    if evento_limpar:
        caminho_modelo_word = ""
        atualizar_label("word", None, None) # Limpa o texto da tela
        return

    arquivo = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
    if arquivo:
        caminho_modelo_word = arquivo
        atualizar_label("word", encurtar_caminho(arquivo), "#2196F3")

def gerenciar_pasta(evento_limpar=False):
    global pasta_destino
    
    if evento_limpar:
        pasta_destino = ""
        atualizar_label("pasta", None, None) # Limpa o texto da tela
        return

    pasta = filedialog.askdirectory()
    if pasta:
        pasta_destino = pasta
        nome_pasta = os.path.basename(pasta)
        atualizar_label("pasta", f"Pasta: {nome_pasta}", "#FF9800")

def formatar_moeda(valor):
    try:
        texto = f"{float(valor):,.2f}"
        return "R$ " + texto.replace(",", "X").replace(".", ",").replace("X", ".")
    except:
        return valor

def executar_automacao():
    if not caminho_excel or not caminho_modelo_word or not pasta_destino:
        messagebox.showwarning("Faltam dados", "Preencha todos os campos antes de gerar!")
        return

    try:
        df = pd.read_excel(caminho_excel, sheet_name="Matriz_Aceitacao")
        arquivos_temp = []
        
        for index, linha in df.iterrows():
            doc = DocxTemplate(caminho_modelo_word)
            contexto = {
                'nome_empresa': linha['Nome da Empresa'],
                'atividade':    linha['Atividade da Empresa'],
                'funcionarios': linha['Funcionários'],
                'gasto_anual':  formatar_moeda(linha['Gasto Anual']),
                'faturamento':  formatar_moeda(linha['Faturamento Anual'])
            }
            doc.render(contexto)
            nome_temp = os.path.join(pasta_destino, f"temp_{index}.docx")
            doc.save(nome_temp)
            arquivos_temp.append(nome_temp)

        if arquivos_temp:
            master = Document(arquivos_temp[0])
            composer = Composer(master)
            for arq in arquivos_temp[1:]:
                master.add_paragraph('\n')
                composer.append(Document(arq))
            
            final = os.path.join(pasta_destino, "Relatorio_Final.docx")
            composer.save(final)
            
            for f in arquivos_temp:
                try: os.remove(f)
                except: pass
            
            messagebox.showinfo("Sucesso", "Relatório Gerado com Sucesso!")
            
    except Exception as e:
        messagebox.showerror("Erro", str(e))

# --- INTERFACE GRÁFICA ---

window = Tk()
window.geometry("598x545")
window.configure(bg = "#FFFFFF")
window.title("Automação de Contratos")

canvas = Canvas(
    window,
    bg = "#FFFFFF",
    height = 545,
    width = 598,
    bd = 0,
    highlightthickness = 0,
    relief = "ridge"
)
canvas.place(x = 0, y = 0)

# Função atualizada para suportar Clique Direito (Apagar)
def create_button(x, y, image_path, command_left, command_right):
    img = PhotoImage(file=relative_to_assets(image_path))
    btn_id = canvas.create_image(x, y, image=img, anchor="nw")
    
    def on_press(event): canvas.move(btn_id, 2, 2)
    
    def on_release_left(event):
        canvas.move(btn_id, -2, -2)
        if command_left: command_left() 

    def on_release_right(event):
        if command_right: command_right() 

    canvas.tag_bind(btn_id, "<Button-1>", on_press)
    canvas.tag_bind(btn_id, "<ButtonRelease-1>", on_release_left)
    canvas.tag_bind(btn_id, "<Button-3>", on_release_right) 
    
    return img

# Fundo
image_image_1 = PhotoImage(file=relative_to_assets("image_1.png"))
canvas.create_image(303.0, 272.0, image=image_image_1)

# --- CRIAÇÃO DOS BOTÕES ---

# Botão Excel
img_excel = create_button(
    443.0, 211.0, "button_4.png", 
    command_left=lambda: gerenciar_excel(False),
    command_right=lambda: gerenciar_excel(True)
)
# REMOVIDO: atualizar_label("excel", ...)

# Botão Word
img_word = create_button(
    443.0, 280.0, "button_3.png", 
    command_left=lambda: gerenciar_word(False),
    command_right=lambda: gerenciar_word(True)
)
# REMOVIDO: atualizar_label("word", ...)

# Botão Pasta
img_pasta = create_button(
    443.0, 349.0, "button_2.png", 
    command_left=lambda: gerenciar_pasta(False),
    command_right=lambda: gerenciar_pasta(True)
)
# REMOVIDO: atualizar_label("pasta", ...)

# Botão Gerar
img_run = create_button(
    249.498, 437.0, "button_1.png", 
    command_left=executar_automacao,
    command_right=None
)

# Retângulos Decorativos
canvas.create_rectangle(12.0, 0.0, 29.0, 15.0, fill="#E6EBEE", outline="")
canvas.create_rectangle(575.0, 0.0, 594.0, 542.0, fill="#E4E5EA", outline="")
canvas.create_rectangle(12.0, 0.0, 25.0, 545.0, fill="#E6E8EE", outline="")
canvas.create_rectangle(12.0, 535.0, 594.0, 547.0, fill="#E4E5EA", outline="")

window.resizable(False, False)
window.mainloop()