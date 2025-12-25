import customtkinter as ctk
from tkinter import filedialog, messagebox
import pandas as pd
from docxtpl import DocxTemplate
from docxcompose.composer import Composer
from docx import Document
import os

# --- CONFIGURAÇÃO VISUAL (MODERNA) ---
ctk.set_appearance_mode("Dark")  # Modos: "System" (padrão), "Dark", "Light"
ctk.set_default_color_theme("blue")  # Temas: "blue", "green", "dark-blue"

class AppAutomacao(ctk.CTk):
    def __init__(self):
        super().__init__()

        # Configuração da Janela Principal
        self.title("Gerador de Relatórios 2.0")
        self.geometry("600x500")
        self.resizable(False, False)

        # Variáveis para guardar os caminhos
        self.caminho_excel = ""
        self.caminho_modelo = ""
        self.pasta_destino = ""

        # --- CRIAÇÃO DOS ELEMENTOS NA TELA ---
        
        # Título
        self.lbl_titulo = ctk.CTkLabel(self, text="Automação de Contratos", font=("Roboto", 24, "bold"))
        self.lbl_titulo.pack(pady=20)

        # 1. Botão Excel
        self.btn_excel = ctk.CTkButton(self, text="1. Selecionar Excel (Benchmark)", command=self.selecionar_excel, height=40)
        self.btn_excel.pack(pady=10, padx=50, fill="x")
        self.lbl_excel = ctk.CTkLabel(self, text="Nenhum arquivo selecionado", text_color="gray")
        self.lbl_excel.pack()

        # 2. Botão Word
        self.btn_word = ctk.CTkButton(self, text="2. Selecionar Modelo Word", command=self.selecionar_word, height=40, fg_color="#D35400", hover_color="#A04000")
        self.btn_word.pack(pady=10, padx=50, fill="x")
        self.lbl_word = ctk.CTkLabel(self, text="Nenhum arquivo selecionado", text_color="gray")
        self.lbl_word.pack()

        # 3. Botão Pasta
        self.btn_pasta = ctk.CTkButton(self, text="3. Onde Salvar?", command=self.selecionar_pasta, height=40, fg_color="#27AE60", hover_color="#1E8449")
        self.btn_pasta.pack(pady=10, padx=50, fill="x")
        self.lbl_pasta = ctk.CTkLabel(self, text="Nenhuma pasta selecionada", text_color="gray")
        self.lbl_pasta.pack()

        # Linha divisória
        self.linha = ctk.CTkProgressBar(self, height=2)
        self.linha.set(0) # Começa vazia
        self.linha.pack(pady=20, padx=50, fill="x")

        # 4. Botão INICIAR (Gigante)
        self.btn_rodar = ctk.CTkButton(self, text="GERAR RELATÓRIO COMPLETO", command=self.rodar_automacao, height=50, font=("Roboto", 16, "bold"), state="disabled")
        self.btn_rodar.pack(pady=10, padx=50, fill="x")

        # Label de Status (Log)
        self.lbl_status = ctk.CTkLabel(self, text="Aguardando dados...", font=("Roboto", 12))
        self.lbl_status.pack(pady=10)

    # --- FUNÇÕES DOS BOTÕES ---
    
    def selecionar_excel(self):
        arquivo = filedialog.askopenfilename(filetypes=[("Excel", "*.xlsx")])
        if arquivo:
            self.caminho_excel = arquivo
            self.lbl_excel.configure(text=f"...{arquivo[-40:]}", text_color="#3498DB") # Mostra só o final do nome
            self.checar_pronto()

    def selecionar_word(self):
        arquivo = filedialog.askopenfilename(filetypes=[("Word", "*.docx")])
        if arquivo:
            self.caminho_modelo = arquivo
            self.lbl_word.configure(text=f"...{arquivo[-40:]}", text_color="#D35400")
            self.checar_pronto()

    def selecionar_pasta(self):
        pasta = filedialog.askdirectory()
        if pasta:
            self.pasta_destino = pasta
            self.lbl_pasta.configure(text=f"...{pasta[-40:]}", text_color="#27AE60")
            self.checar_pronto()

    def checar_pronto(self):
        # Só libera o botão de rodar se tiver tudo preenchido
        if self.caminho_excel and self.caminho_modelo and self.pasta_destino:
            self.btn_rodar.configure(state="normal", fg_color="#1F618D")

    def formatar_moeda(self, valor):
        try:
            texto = f"{float(valor):,.2f}"
            return "R$ " + texto.replace(",", "X").replace(".", ",").replace("X", ".")
        except:
            return valor

    # --- O CÉREBRO DA AUTOMAÇÃO ---
    def rodar_automacao(self):
        self.btn_rodar.configure(state="disabled", text="PROCESSANDO...")
        self.linha.set(0)
        self.lbl_status.configure(text="Lendo Excel...")
        self.update() # Atualiza a tela para não travar

        try:
            # 1. Leitura
            df = pd.read_excel(self.caminho_excel, sheet_name="Matriz_Aceitacao")
            total = len(df)
            
            arquivos_temporarios = []
            
            # 2. Processamento
            for index, linha in df.iterrows():
                # Atualiza a barra de progresso visualmente
                progresso = (index + 1) / total
                self.linha.set(progresso)
                self.lbl_status.configure(text=f"Processando: {linha['Nome da Empresa']}")
                self.update()

                doc = DocxTemplate(self.caminho_modelo)
                contexto = {
                    'nome_empresa': linha['Nome da Empresa'],
                    'atividade':    linha['Atividade da Empresa'],
                    'funcionarios': linha['Funcionários'],
                    'gasto_anual':  self.formatar_moeda(linha['Gasto Anual']),
                    'faturamento':  self.formatar_moeda(linha['Faturamento Anual'])
                }
                doc.render(contexto)
                
                nome_temp = os.path.join(self.pasta_destino, f"temp_{index}.docx")
                doc.save(nome_temp)
                arquivos_temporarios.append(nome_temp)

            # 3. Consolidação
            self.lbl_status.configure(text="Unindo arquivos (isso pode demorar um pouco)...")
            self.update()

            if arquivos_temporarios:
                master_doc = Document(arquivos_temporarios[0])
                composer = Composer(master_doc)
                
                for arq_temp in arquivos_temporarios[1:]:
                    doc_to_append = Document(arq_temp)
                    master_doc.add_paragraph('\n') 
                    composer.append(doc_to_append)
                
                caminho_final = os.path.join(self.pasta_destino, "Relatorio_Final_Completo.docx")
                composer.save(caminho_final)
                
                # Limpeza
                for f in arquivos_temporarios:
                    try:
                        os.remove(f)
                    except:
                        pass

                self.lbl_status.configure(text="CONCLUÍDO COM SUCESSO!")
                self.linha.set(1)
                messagebox.showinfo("Sucesso", f"Arquivo gerado em:\n{caminho_final}")
            
            else:
                messagebox.showwarning("Aviso", "O Excel estava vazio ou não foi lido corretamente.")

        except Exception as e:
            messagebox.showerror("Erro Crítico", f"Ocorreu um erro:\n{e}")
            self.lbl_status.configure(text="Erro fatal.")
        
        # Restaura o botão
        self.btn_rodar.configure(state="normal", text="GERAR RELATÓRIO COMPLETO")

# --- INICIALIZAR O APP ---
if __name__ == "__main__":
    app = AppAutomacao()
    app.mainloop()