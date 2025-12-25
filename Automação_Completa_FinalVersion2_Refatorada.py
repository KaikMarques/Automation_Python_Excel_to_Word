import sys
import os
import threading # ADICIONADO: Para multitarefa
from pathlib import Path
from tkinter import Tk, Canvas, PhotoImage, filedialog, messagebox, Toplevel # ADICIONADO: Toplevel
from tkinter import ttk # ADICIONADO: Para barra de progresso
import pandas as pd
from docxtpl import DocxTemplate
from docxcompose.composer import Composer
from docx import Document

# --- LÓGICA DE CAMINHOS: VS CODE vs EXECUTÁVEL ---
if getattr(sys, 'frozen', False):
    base_path = sys._MEIPASS
    ASSETS_PATH = Path(base_path) / "assets" / "frame0"
else:
    # Ajuste este caminho se necessário
    ASSETS_PATH = Path(r"C:\Users\kaikm\OneDrive\Documentos\Projetos\Automação\Automation_Python_Excel_to_Word\design\build\assets\frame0")

def relative_to_assets(path: str) -> Path:
    return ASSETS_PATH / Path(path)

# --- VARIÁVEIS GLOBAIS ---
caminho_excel = ""
caminho_modelo_word = ""
caminho_saida_final = ""

# Dicionário para guardar os IDs dos textos na tela
elementos_texto = {
    "excel": None,
    "word": None,
    "pasta": None
}

# --- FUNÇÕES DE AUXÍLIO ---

def encurtar_caminho(caminho):
    if not caminho: return ""
    nome = os.path.basename(caminho)
    if len(nome) > 30: return "..." + nome[-27:]
    return nome

def atualizar_label(tipo, texto, cor):
    if elementos_texto[tipo]:
        canvas.delete(elementos_texto[tipo])
    
    if not texto: return

    posicoes_y = {"excel": 222, "word": 291, "pasta": 360}
    
    id_novo = canvas.create_text(
        430.0, posicoes_y[tipo], text=texto, anchor="e",
        fill=cor, font=("Arial", 10, "bold")
    )
    elementos_texto[tipo] = id_novo

def formatar_moeda(valor):
    try:
        texto = f"{float(valor):,.2f}"
        return "R$ " + texto.replace(",", "X").replace(".", ",").replace("X", ".")
    except:
        return valor

# --- FUNÇÕES DOS BOTÕES (LÓGICA) ---

def gerenciar_excel(evento_limpar=False):
    global caminho_excel
    if evento_limpar:
        caminho_excel = ""
        atualizar_label("excel", None, None)
        return
    arquivo = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx")])
    if arquivo:
        caminho_excel = arquivo
        atualizar_label("excel", encurtar_caminho(arquivo), "#4CAF50") 

def gerenciar_word(evento_limpar=False):
    global caminho_modelo_word
    if evento_limpar:
        caminho_modelo_word = ""
        atualizar_label("word", None, None)
        return
    arquivo = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
    if arquivo:
        caminho_modelo_word = arquivo
        atualizar_label("word", encurtar_caminho(arquivo), "#2196F3")

def gerenciar_saida(evento_limpar=False):
    global caminho_saida_final
    if evento_limpar:
        caminho_saida_final = ""
        atualizar_label("pasta", None, None)
        return
    
    arquivo_saida = filedialog.asksaveasfilename(
        parent=window,
        title="Salvar Relatório Como...",
        defaultextension=".docx",
        filetypes=[("Documento Word", "*.docx")]
    )
    
    if arquivo_saida:
        caminho_saida_final = arquivo_saida
        nome_arquivo = os.path.basename(arquivo_saida)
        atualizar_label("pasta", f"Saída: {nome_arquivo}", "#FF9800")

# --- LÓGICA DE AUTOMACAO (THREADING + PROGRESSO) ---

def worker_automacao(janela_progresso, barra_progresso, label_progresso):
    """Esta função roda em segundo plano para não travar a tela"""
    try:
        pasta_base = os.path.dirname(caminho_saida_final)
        
        # 1. Leitura do Excel
        label_progresso.config(text="Lendo Excel...")
        df = pd.read_excel(caminho_excel, sheet_name="Matriz_Aceitacao")
        total_linhas = len(df)
        arquivos_temp = []
        
        # Configura a barra: total de linhas + 20% extra para a fusão
        barra_progresso["maximum"] = total_linhas * 1.2 
        
        # 2. Geração dos Arquivos Individuais
        for index, linha in df.iterrows():
            label_progresso.config(text=f"Gerando item {index + 1} de {total_linhas}...")
            
            doc = DocxTemplate(caminho_modelo_word)
            contexto = {
                'nome_empresa': linha['Nome da Empresa'],
                'atividade':    linha['Atividade da Empresa'],
                'funcionarios': linha['Funcionários'],
                'gasto_anual':  formatar_moeda(linha['Gasto Anual']),
                'faturamento':  formatar_moeda(linha['Faturamento Anual'])
            }
            doc.render(contexto)
            
            # Salva temporário
            nome_temp = os.path.join(pasta_base, f"temp_{index}.docx")
            doc.save(nome_temp)
            arquivos_temp.append(nome_temp)
            
            # Atualiza barra
            barra_progresso["value"] = index + 1
        
        # 3. Unificação (Merge)
        if arquivos_temp:
            label_progresso.config(text="Unificando arquivos... (Aguarde)")
            master = Document(arquivos_temp[0])
            composer = Composer(master)
            
            # Etapa de merge (pode ser lenta com 300 arquivos)
            qtd_temp = len(arquivos_temp)
            for i, arq in enumerate(arquivos_temp[1:]):
                master.add_paragraph('\n')
                composer.append(Document(arq))
                
                # Incremento visual pequeno durante a fusão
                progresso_atual = total_linhas + ((i / qtd_temp) * (total_linhas * 0.2))
                barra_progresso["value"] = progresso_atual
            
            label_progresso.config(text="Salvando arquivo final no disco...")
            composer.save(caminho_saida_final)
            
            # 4. Limpeza
            label_progresso.config(text="Limpando arquivos temporários...")
            for f in arquivos_temp:
                try: os.remove(f)
                except: pass
            
            # Fecha a janelinha e mostra sucesso
            janela_progresso.destroy()
            messagebox.showinfo("Concluído", f"Sucesso!\nArquivo salvo em:\n{encurtar_caminho(caminho_saida_final)}")
            
    except Exception as e:
        janela_progresso.destroy()
        messagebox.showerror("Erro no Processo", str(e))

def executar_automacao():
    if not caminho_excel or not caminho_modelo_word or not caminho_saida_final:
        messagebox.showwarning("Atenção", "Preencha Excel, Modelo e Local de Saída antes de rodar.")
        return

    # Cria janela Pop-up de Progresso
    janela_progresso = Toplevel(window)
    janela_progresso.title("Processando...")
    janela_progresso.geometry("350x150")
    janela_progresso.resizable(False, False)
    
    # Estiliza para centralizar na tela (opcional, mas fica melhor)
    x_c = window.winfo_x() + 120
    y_c = window.winfo_y() + 150
    janela_progresso.geometry(f"+{x_c}+{y_c}")
    
    # Texto de status
    lbl_status = ttk.Label(janela_progresso, text="Iniciando automação...", font=("Arial", 10))
    lbl_status.pack(pady=(30, 10))
    
    # Barra de Progresso
    progress = ttk.Progressbar(janela_progresso, orient="horizontal", length=280, mode="determinate")
    progress.pack(pady=10)
    
    # Inicia a Thread (Processamento Paralelo)
    t = threading.Thread(target=worker_automacao, args=(janela_progresso, progress, lbl_status))
    t.start()

# --- INTERFACE GRÁFICA ---

window = Tk()
window.geometry("598x545")
window.configure(bg = "#FFFFFF")
window.title("Automação de Contratos - v2.0") # Mudei o título para v2.0

canvas = Canvas(
    window,
    bg = "#FFFFFF",
    height = 545,
    width = 598,
    bd = 0,
    highlightthickness = 0,
    relief = "ridge"
)
canvas.place(x = 0, y = 0)

def create_button(x, y, image_path, command_left, command_right):
    img = PhotoImage(file=relative_to_assets(image_path))
    btn_id = canvas.create_image(x, y, image=img, anchor="nw")
    
    def on_press(event): canvas.move(btn_id, 2, 2)
    def on_release_left(event):
        canvas.move(btn_id, -2, -2)
        if command_left: command_left() 
    def on_release_right(event):
        if command_right: command_right() 

    canvas.tag_bind(btn_id, "<Button-1>", on_press)
    canvas.tag_bind(btn_id, "<ButtonRelease-1>", on_release_left)
    canvas.tag_bind(btn_id, "<Button-3>", on_release_right) 
    return img

image_image_1 = PhotoImage(file=relative_to_assets("image_1.png"))
canvas.create_image(303.0, 272.0, image=image_image_1)

# Botões
img_excel = create_button(443.0, 211.0, "button_4.png", lambda: gerenciar_excel(False), lambda: gerenciar_excel(True))
img_word = create_button(443.0, 280.0, "button_3.png", lambda: gerenciar_word(False), lambda: gerenciar_word(True))
img_pasta = create_button(443.0, 349.0, "button_2.png", lambda: gerenciar_saida(False), lambda: gerenciar_saida(True))
img_run = create_button(249.498, 437.0, "button_1.png", executar_automacao, None)

# Decorações
canvas.create_rectangle(12.0, 0.0, 29.0, 15.0, fill="#E6EBEE", outline="")
canvas.create_rectangle(575.0, 0.0, 594.0, 542.0, fill="#E4E5EA", outline="")
canvas.create_rectangle(12.0, 0.0, 25.0, 545.0, fill="#E6E8EE", outline="")
canvas.create_rectangle(12.0, 535.0, 594.0, 547.0, fill="#E4E5EA", outline="")

window.resizable(False, False)
window.mainloop()