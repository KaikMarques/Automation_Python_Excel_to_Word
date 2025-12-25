import sys
import os
from pathlib import Path
from tkinter import Tk, Canvas, PhotoImage, filedialog, messagebox
import pandas as pd
from docxtpl import DocxTemplate
from docxcompose.composer import Composer
from docx import Document

# --- LÓGICA DE CAMINHOS: VS CODE vs EXECUTÁVEL ---
if getattr(sys, 'frozen', False):
    base_path = sys._MEIPASS
    ASSETS_PATH = Path(base_path) / "assets" / "frame0"
else:
    ASSETS_PATH = Path(r"C:\Users\kaikm\OneDrive\Documentos\Projetos\Automação\Automation_Python_Excel_to_Word\design\build\assets\frame0")

def relative_to_assets(path: str) -> Path:
    return ASSETS_PATH / Path(path)

# --- VARIÁVEIS GLOBAIS ---
caminho_excel = ""
caminho_modelo_word = ""
caminho_saida_final = "" # Mudei o nome da variável para fazer mais sentido

# Dicionário para guardar os IDs dos textos na tela
elementos_texto = {
    "excel": None,
    "word": None,
    "pasta": None
}

# --- FUNÇÕES DE AUXÍLIO ---

def encurtar_caminho(caminho):
    if not caminho: return ""
    nome = os.path.basename(caminho)
    if len(nome) > 30: return "..." + nome[-27:]
    return nome

def atualizar_label(tipo, texto, cor):
    if elementos_texto[tipo]:
        canvas.delete(elementos_texto[tipo])
    
    if not texto: return

    posicoes_y = {"excel": 222, "word": 291, "pasta": 360}
    
    id_novo = canvas.create_text(
        430.0, posicoes_y[tipo], text=texto, anchor="e",
        fill=cor, font=("Arial", 10, "bold")
    )
    elementos_texto[tipo] = id_novo

# --- FUNÇÕES DOS BOTÕES (LÓGICA) ---

def gerenciar_excel(evento_limpar=False):
    global caminho_excel
    if evento_limpar:
        caminho_excel = ""
        atualizar_label("excel", None, None)
        return
    arquivo = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx")])
    if arquivo:
        caminho_excel = arquivo
        atualizar_label("excel", encurtar_caminho(arquivo), "#4CAF50") 

def gerenciar_word(evento_limpar=False):
    global caminho_modelo_word
    if evento_limpar:
        caminho_modelo_word = ""
        atualizar_label("word", None, None)
        return
    arquivo = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
    if arquivo:
        caminho_modelo_word = arquivo
        atualizar_label("word", encurtar_caminho(arquivo), "#2196F3")

# --- MUDANÇA PRINCIPAL AQUI (Salvar Como) ---
def gerenciar_saida(evento_limpar=False):
    global caminho_saida_final
    if evento_limpar:
        caminho_saida_final = ""
        atualizar_label("pasta", None, None)
        return
    
    # Abre janela "Salvar Como"
    arquivo_saida = filedialog.asksaveasfilename(
        parent=window,
        title="Salvar Relatório Como...",
        defaultextension=".docx",
        filetypes=[("Documento Word", "*.docx")]
    )
    
    if arquivo_saida:
        caminho_saida_final = arquivo_saida
        nome_arquivo = os.path.basename(arquivo_saida)
        # Mostra o nome do arquivo que será criado
        atualizar_label("pasta", f"Saída: {nome_arquivo}", "#FF9800")

def formatar_moeda(valor):
    try:
        texto = f"{float(valor):,.2f}"
        return "R$ " + texto.replace(",", "X").replace(".", ",").replace("X", ".")
    except:
        return valor

def executar_automacao():
    if not caminho_excel or not caminho_modelo_word or not caminho_saida_final:
        messagebox.showwarning("Faltam dados", "Preencha todos os campos antes de gerar!")
        return

    try:
        # Pega a pasta onde o arquivo final será salvo (para colocar os temporários lá)
        pasta_base = os.path.dirname(caminho_saida_final)
        
        df = pd.read_excel(caminho_excel, sheet_name="Matriz_Aceitacao")
        arquivos_temp = []
        
        for index, linha in df.iterrows():
            doc = DocxTemplate(caminho_modelo_word)
            contexto = {
                'nome_empresa': linha['Nome da Empresa'],
                'atividade':    linha['Atividade da Empresa'],
                'funcionarios': linha['Funcionários'],
                'gasto_anual':  formatar_moeda(linha['Gasto Anual']),
                'faturamento':  formatar_moeda(linha['Faturamento Anual'])
            }
            doc.render(contexto)
            
            # Salva o temporário na mesma pasta do arquivo final
            nome_temp = os.path.join(pasta_base, f"temp_{index}.docx")
            doc.save(nome_temp)
            arquivos_temp.append(nome_temp)

        if arquivos_temp:
            master = Document(arquivos_temp[0])
            composer = Composer(master)
            for arq in arquivos_temp[1:]:
                master.add_paragraph('\n')
                composer.append(Document(arq))
            
            # Salva com o nome EXATO que o usuário escolheu
            composer.save(caminho_saida_final)
            
            # Limpa temporários
            for f in arquivos_temp:
                try: os.remove(f)
                except: pass
            
            messagebox.showinfo("Sucesso", f"Relatório salvo em:\n{caminho_saida_final}")
            
    except Exception as e:
        messagebox.showerror("Erro", str(e))

# --- INTERFACE GRÁFICA ---

window = Tk()
window.geometry("598x545")
window.configure(bg = "#FFFFFF")
window.title("Automação de Contratos")

canvas = Canvas(
    window,
    bg = "#FFFFFF",
    height = 545,
    width = 598,
    bd = 0,
    highlightthickness = 0,
    relief = "ridge"
)
canvas.place(x = 0, y = 0)

def create_button(x, y, image_path, command_left, command_right):
    img = PhotoImage(file=relative_to_assets(image_path))
    btn_id = canvas.create_image(x, y, image=img, anchor="nw")
    
    def on_press(event): canvas.move(btn_id, 2, 2)
    def on_release_left(event):
        canvas.move(btn_id, -2, -2)
        if command_left: command_left() 
    def on_release_right(event):
        if command_right: command_right() 

    canvas.tag_bind(btn_id, "<Button-1>", on_press)
    canvas.tag_bind(btn_id, "<ButtonRelease-1>", on_release_left)
    canvas.tag_bind(btn_id, "<Button-3>", on_release_right) 
    return img

image_image_1 = PhotoImage(file=relative_to_assets("image_1.png"))
canvas.create_image(303.0, 272.0, image=image_image_1)

# Botões
img_excel = create_button(443.0, 211.0, "button_4.png", lambda: gerenciar_excel(False), lambda: gerenciar_excel(True))
img_word = create_button(443.0, 280.0, "button_3.png", lambda: gerenciar_word(False), lambda: gerenciar_word(True))

# Atualizado para chamar gerenciar_saida
img_pasta = create_button(443.0, 349.0, "button_2.png", lambda: gerenciar_saida(False), lambda: gerenciar_saida(True))

img_run = create_button(249.498, 437.0, "button_1.png", executar_automacao, None)

canvas.create_rectangle(12.0, 0.0, 29.0, 15.0, fill="#E6EBEE", outline="")
canvas.create_rectangle(575.0, 0.0, 594.0, 542.0, fill="#E4E5EA", outline="")
canvas.create_rectangle(12.0, 0.0, 25.0, 545.0, fill="#E6E8EE", outline="")
canvas.create_rectangle(12.0, 535.0, 594.0, 547.0, fill="#E4E5EA", outline="")

window.resizable(False, False)
window.mainloop()